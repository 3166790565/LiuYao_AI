# utils/export.py
# 导出功能工具函数

import datetime
import os
from typing import Dict, List, Tuple

# 导入配置
from config.constants import APP_NAME, APP_VERSION
from config.constants import FILE_TYPES, TEXT
from utils.logger import setup_logger, log_exception

# 设置日志记录器
logger = setup_logger(__name__)


def export_to_pdf(
    file_path: str,
    title: str,
    question: str,
    hexagram_info: str,
    analysis_results: List[Dict[str, str]],
    conclusion: str,
    language: str = "zh_CN"
) -> Tuple[bool, str]:
    """导出分析结果为PDF文件
    
    Args:
        file_path: 文件保存路径
        title: 标题
        question: 问题
        hexagram_info: 卦象信息
        analysis_results: 分析结果列表
        conclusion: 总结论
        language: 语言代码
        
    Returns:
        Tuple[bool, str]: (成功标志, 成功/错误消息)
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
    except ImportError:
        error_msg = TEXT[language]["pdf_library_missing"]
        logger.error(error_msg)
        return False, error_msg
    
    try:
        # 注册中文字体
        try:
            # 尝试注册微软雅黑字体
            pdfmetrics.registerFont(TTFont('Microsoft YaHei', 'msyh.ttc'))
            font_name = 'Microsoft YaHei'
        except:
            # 如果失败，使用内置字体
            font_name = 'Helvetica'
        
        # 创建PDF文档
        doc = SimpleDocTemplate(
            file_path,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # 创建样式
        styles = getSampleStyleSheet()
        
        # 检查并添加自定义标题样式
        if 'CustomTitle' not in styles:
            styles.add(ParagraphStyle(
                name='CustomTitle',
                fontName=font_name,
                fontSize=18,
                alignment=TA_CENTER,
                spaceAfter=12
            ))
        
        # 检查并添加自定义标题样式
        if 'CustomHeading' not in styles:
            styles.add(ParagraphStyle(
                name='CustomHeading',
                fontName=font_name,
                fontSize=14,
                alignment=TA_LEFT,
                spaceAfter=6
            ))
        # 检查并添加自定义正文样式
        if 'CustomNormal' not in styles:
            styles.add(ParagraphStyle(
                name='CustomNormal',
                fontName=font_name,
                fontSize=10,
                alignment=TA_LEFT,
                spaceAfter=6
            ))
        
        # 创建内容
        content = []
        
        # 标题
        content.append(Paragraph(title, styles['CustomTitle']))
        content.append(Spacer(1, 12))
        
        # 问题
        content.append(Paragraph(f"{TEXT[language]['question']}: {question}", styles['CustomHeading']))
        content.append(Spacer(1, 6))
        
        # 卦象信息
        content.append(Paragraph(f"{TEXT[language]['hexagram_info']}:", styles['CustomHeading']))
        content.append(Paragraph(hexagram_info, styles['CustomNormal']))
        content.append(Spacer(1, 12))
        
        # 分析结果
        content.append(Paragraph(f"{TEXT[language]['analysis_result']}:", styles['CustomHeading']))
        
        # 处理分析结果数据
        if isinstance(analysis_results, str):
            # 如果是字符串，直接添加
            content.append(Paragraph(analysis_results, styles['CustomNormal']))
        elif isinstance(analysis_results, list):
            # 如果是列表，按原逻辑处理
            for result in analysis_results:
                if isinstance(result, dict):
                    aspect = result.get('aspect', '')
                    interpretation = result.get('interpretation', '')
                    content.append(Paragraph(f"<b>{aspect}</b>", styles['CustomNormal']))
                    content.append(Paragraph(interpretation, styles['CustomNormal']))
                else:
                    # 如果列表元素是字符串
                    content.append(Paragraph(str(result), styles['CustomNormal']))
                content.append(Spacer(1, 6))
        else:
            # 其他类型转为字符串
            content.append(Paragraph(str(analysis_results), styles['CustomNormal']))
        
        content.append(Spacer(1, 6))
        
        # 总结论
        content.append(Paragraph(f"{TEXT[language]['conclusion']}:", styles['CustomHeading']))
        content.append(Paragraph(conclusion, styles['CustomNormal']))
        content.append(Spacer(1, 12))
        
        # 导出时间
        export_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        content.append(Paragraph(f"{TEXT[language]['export_time']}: {export_time}", styles['CustomNormal']))
        
        # 版本信息
        content.append(Paragraph(f"{APP_NAME} {TEXT[language]['version']} {APP_VERSION}", styles['CustomNormal']))
        
        # 构建PDF
        doc.build(content)
        
        return True, file_path
    except Exception as e:
        error_msg = f"{TEXT[language]['pdf_export_error']}: {str(e)}"
        log_exception(e, logger)
        return False, error_msg


def export_to_word(
    file_path: str,
    title: str,
    question: str,
    hexagram_info: str,
    analysis_results: List[Dict[str, str]],
    conclusion: str,
    language: str = "zh_CN"
) -> Tuple[bool, str]:
    """导出分析结果为Word文档
    
    Args:
        file_path: 文件保存路径
        title: 标题
        question: 问题
        hexagram_info: 卦象信息
        analysis_results: 分析结果列表
        conclusion: 总结论
        language: 语言代码
        
    Returns:
        Tuple[bool, str]: (成功标志, 成功/错误消息)
    """
    try:
        import docx
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        error_msg = TEXT[language]["word_library_missing"]
        logger.error(error_msg)
        return False, error_msg
    
    try:
        # 创建Word文档
        doc = docx.Document()
        
        # 设置标题
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        
        # 添加问题
        question_para = doc.add_paragraph()
        question_para.add_run(f"{TEXT[language]['question']}: ").font.bold = True
        question_para.add_run(question)
        
        # 添加卦象信息
        doc.add_heading(TEXT[language]['hexagram_info'], level=2)
        doc.add_paragraph(hexagram_info)
        
        # 添加分析结果
        doc.add_heading(TEXT[language]['analysis_result'], level=2)
        
        # 处理分析结果数据
        if isinstance(analysis_results, str):
            # 如果是字符串，直接添加
            doc.add_paragraph(analysis_results)
        elif isinstance(analysis_results, list):
            # 如果是列表，按原逻辑处理
            for result in analysis_results:
                if isinstance(result, dict):
                    aspect = result.get('aspect', '')
                    interpretation = result.get('interpretation', '')
                    
                    aspect_para = doc.add_paragraph()
                    aspect_run = aspect_para.add_run(aspect)
                    aspect_run.font.bold = True
                    
                    doc.add_paragraph(interpretation)
                else:
                    # 如果列表元素是字符串
                    doc.add_paragraph(str(result))
        else:
            # 其他类型转为字符串
            doc.add_paragraph(str(analysis_results))
        
        # 添加总结论
        doc.add_heading(TEXT[language]['conclusion'], level=2)
        doc.add_paragraph(conclusion)
        
        # 添加导出时间
        export_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        time_para = doc.add_paragraph()
        time_para.add_run(f"{TEXT[language]['export_time']}: {export_time}")
        
        # 添加版本信息
        version_para = doc.add_paragraph()
        version_para.add_run(f"{APP_NAME} {TEXT[language]['version']} {APP_VERSION}")
        
        # 保存文档
        doc.save(file_path)
        
        return True, file_path
    except Exception as e:
        error_msg = f"{TEXT[language]['word_export_error']}: {str(e)}"
        log_exception(e, logger)
        return False, error_msg


def export_to_text(
    file_path: str,
    title: str,
    question: str,
    hexagram_info: str,
    analysis_results: List[Dict[str, str]],
    conclusion: str,
    language: str = "zh_CN"
) -> Tuple[bool, str]:
    """导出分析结果为文本文件
    
    Args:
        file_path: 文件保存路径
        title: 标题
        question: 问题
        hexagram_info: 卦象信息
        analysis_results: 分析结果列表
        conclusion: 总结论
        language: 语言代码
        
    Returns:
        Tuple[bool, str]: (成功标志, 成功/错误消息)
    """
    try:
        with open(file_path, 'w', encoding='utf-8') as f:
            # 写入标题
            f.write(f"{title}\n\n")
            
            # 写入问题
            f.write(f"{TEXT[language]['question']}: {question}\n\n")
            
            # 写入卦象信息
            f.write(f"{TEXT[language]['hexagram_info']}:\n")
            f.write(f"{hexagram_info}\n\n")
            
            # 写入分析结果
            f.write(f"{TEXT[language]['analysis_result']}:\n")
            
            # 处理分析结果数据
            if isinstance(analysis_results, str):
                # 如果是字符串，直接写入
                f.write(f"{analysis_results}\n\n")
            elif isinstance(analysis_results, list):
                # 如果是列表，按原逻辑处理
                for result in analysis_results:
                    if isinstance(result, dict):
                        aspect = result.get('aspect', '')
                        interpretation = result.get('interpretation', '')
                        f.write(f"{aspect}\n")
                        f.write(f"{interpretation}\n\n")
                    else:
                        # 如果列表元素是字符串
                        f.write(f"{str(result)}\n\n")
            else:
                # 其他类型转为字符串
                f.write(f"{str(analysis_results)}\n\n")
            
            # 写入总结论
            f.write(f"{TEXT[language]['conclusion']}:\n")
            f.write(f"{conclusion}\n\n")
            
            # 写入导出时间
            export_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            f.write(f"{TEXT[language]['export_time']}: {export_time}\n")
            
            # 写入版本信息
            f.write(f"{APP_NAME} {TEXT[language]['version']} {APP_VERSION}\n")
        
        return True, file_path
    except Exception as e:
        error_msg = f"导出文本文件错误: {str(e)}"
        log_exception(e, logger)
        return False, error_msg


def export_analysis_results(
    file_path: str,
    title: str,
    question: str,
    hexagram_info: str,
    analysis_results: List[Dict[str, str]],
    conclusion: str,
    language: str = "zh_CN"
) -> Tuple[bool, str]:
    """根据文件类型导出分析结果
    
    Args:
        file_path: 文件保存路径
        title: 标题
        question: 问题
        hexagram_info: 卦象信息
        analysis_results: 分析结果列表
        conclusion: 总结论
        language: 语言代码
        
    Returns:
        Tuple[bool, str]: (成功标志, 成功/错误消息)
    """
    # 获取文件扩展名
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    # 根据文件类型调用相应的导出函数
    if ext == FILE_TYPES["PDF"]["extension"]:
        return export_to_pdf(file_path, title, question, hexagram_info, analysis_results, conclusion, language)
    elif ext == FILE_TYPES["WORD"]["extension"]:
        return export_to_word(file_path, title, question, hexagram_info, analysis_results, conclusion, language)
    elif ext == FILE_TYPES["TEXT"]["extension"]:
        return export_to_text(file_path, title, question, hexagram_info, analysis_results, conclusion, language)
    else:
        error_msg = f"不支持的文件类型: {ext}"
        logger.error(error_msg)
        return False, error_msg